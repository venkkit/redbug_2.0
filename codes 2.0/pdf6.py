# %%
from docx import Document
from docx.shared import Inches
from docx.shared import Cm
from docx.shared import Mm
import os

document = Document()

section = document.sections[0]
section.page_height = Mm(297)
section.page_width = Mm(210)
section.left_margin = Cm(1)
section.top_margin = Cm(0.5)
section.right_margin = Cm(0.5)
section.bottom_margin = Cm(0.5)



header = section.header
header_para = header.paragraphs[0]
header_para.text = " .  "
header_para.text = "  . "





para = document.add_paragraph()
para.paragraph_format.line_spacing = Cm(9)




"""
document.add_picture('1.jpg', width=Cm(9))
document.add_picture('2.jpg', width=Cm(9))
document.add_picture('3.jpg', width=Cm(9))
document.add_picture('4.jpg', width=Cm(9))

"""


# %%

list =[]

# %%

# To find the images count
for image_name in os.listdir("output/images"):
    # add image so it creates a row with 4 images
    if image_name.endswith("jpg"):
        #print(image_name)
        list.append(image_name)
        

        # printing file name of desired extension
    else:
        continue


# %%
#print(len(list))
list.sort(key=lambda f: int(''.join(filter(str.isdigit, f))))


# %%
# add images to docx
for x in list:
    print(x)
    r = para.add_run()
    r.add_picture(f"output/images/"+x, width=Cm(9))

   

# %%
document.save('output/pdf_slips6.docx')


